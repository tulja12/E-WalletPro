import base64
import re
import sys
import tempfile
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


LINK_PATTERN = re.compile(r"\[([^\]]+)\]\([^)]+\)")
INLINE_CODE_PATTERN = re.compile(r"`([^`]+)`")
NUMBERED_PATTERN = re.compile(r"^\d+\.\s+")


def clean_inline(text: str) -> str:
    text = LINK_PATTERN.sub(r"\1", text)
    text = INLINE_CODE_PATTERN.sub(r"\1", text)
    return text


def set_run_font(run, name="Calibri", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_plain_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run)
    return paragraph


def render_mermaid_to_file(diagram_text: str) -> Path | None:
    encoded = base64.b64encode(diagram_text.encode("utf-8")).decode("ascii")
    request = urllib.request.Request(
        f"https://mermaid.ink/img/{encoded}",
        headers={"User-Agent": "Mozilla/5.0"},
    )
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            image_bytes = response.read()
    except Exception:
        return None

    temp_dir = Path(tempfile.gettempdir())
    output_path = temp_dir / f"mermaid_{abs(hash(diagram_text))}.jpg"
    output_path.write_bytes(image_bytes)
    return output_path


def add_code_block(document: Document, code_lines: list[str], language: str):
    if language == "mermaid":
        image_path = render_mermaid_to_file("\n".join(code_lines).strip())
        if image_path and image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.2))
            return

    for line in code_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9)


def parse_table_row(line: str) -> list[str]:
    row = line.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [clean_inline(cell.strip()) for cell in row.split("|")]


def is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|").replace(":", "").replace("-", "").replace(" ", "")
    return stripped == ""


def add_table(document: Document, lines: list[str]):
    rows = [parse_table_row(line) for line in lines if not is_table_separator(line)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            value = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, bold=(row_index == 0))


def convert_markdown_to_docx(source: Path, destination: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    document.core_properties.title = "E-Wallet Project SRS"
    document.core_properties.subject = "Software Requirements Specification"

    index = 0
    pending_blank = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped == "":
            pending_blank = True
            index += 1
            continue

        if stripped == "---":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            pending_blank = False
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip().lower()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines, language)
            pending_blank = False
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            pending_blank = False
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            text = clean_inline(stripped[level:].strip())
            document.add_heading(text, level=level if level > 0 else 1)
            pending_blank = False
            index += 1
            continue

        if stripped.startswith("- "):
            add_plain_paragraph(document, stripped[2:].strip(), style="List Bullet")
            pending_blank = False
            index += 1
            continue

        if NUMBERED_PATTERN.match(stripped):
            item_text = NUMBERED_PATTERN.sub("", stripped).strip()
            add_plain_paragraph(document, item_text, style="List Number")
            pending_blank = False
            index += 1
            continue

        if pending_blank:
            document.add_paragraph()
            pending_blank = False

        add_plain_paragraph(document, stripped)
        index += 1

    document.save(destination)


def main():
    if len(sys.argv) != 3:
        print("Usage: python tools/export_srs_docx.py <source-md> <output-docx>")
        sys.exit(1)

    source = Path(sys.argv[1]).resolve()
    destination = Path(sys.argv[2]).resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)
    convert_markdown_to_docx(source, destination)
    print(f"Created DOCX: {destination}")


if __name__ == "__main__":
    main()
